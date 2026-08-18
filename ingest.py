"""
Load timetable files from data/ into a common record list.

To update a semester:
  1. Replace files in data/ (same folder, any filename is fine).
  2. Optional: drop venues.csv in a department folder to override parsing.
  3. Add a new programme in config.json (id, name, icon, parser, path).
"""

from __future__ import annotations

import glob
import os
import re
import openpyxl
import xlrd

DAYS = ["MON", "TUE", "WED", "THU", "FRI"]
PERIOD_TIMES = {
    1: "8:30-9:20",
    2: "9:20-10:10",
    3: "10:20-11:10",
    4: "11:10-12:00",
    5: "12:40-1:30",
    6: "1:30-2:20",
    7: "2:20-3:05",
    8: "3:05-3:50",
}

# Standard SRM merged-column layout (1-indexed): C,E,H,J,M,O,Q,S
SRM_WIDE_COLS = {3: 1, 5: 2, 8: 3, 10: 4, 13: 5, 15: 6, 17: 7, 19: 8}
SKIP_SHEETS = ("sheet", "copy of", "ostin")


def year_sort(y: str) -> int:
    return {"I Year": 1, "II Year": 2, "III Year": 3, "IV Year": 4}.get(y, 99)


def year_from_text(s: str) -> str:
    if not s:
        return ""
    u = s.upper()
    if re.search(r"\bIV\b|4TH|YEAR\s*4|IV\s*YEAR", u):
        return "IV Year"
    if re.search(r"\bIII\b|3RD|YEAR\s*3|III\s*YEAR", u):
        return "III Year"
    if re.search(r"\bII\b|2ND|YEAR\s*2|II\s*YEAR", u):
        return "II Year"
    if re.search(r"\bI\b|1ST|YEAR\s*1|I\s*YEAR", u):
        return "I Year"
    m = re.search(r"(I{1,3}V?)\s*/", s, re.IGNORECASE)
    if m:
        return {"I": "I Year", "II": "II Year", "III": "III Year", "IV": "IV Year"}.get(
            m.group(1).upper(), ""
        )
    return ""


def year_from_info(s: str) -> str:
    if not s:
        return ""
    m = re.search(r"(I{1,3}V?)\s*/", s, re.IGNORECASE)
    if m:
        return {"I": "I Year", "II": "II Year", "III": "III Year", "IV": "IV Year"}.get(
            m.group(1).upper(), ""
        )
    return year_from_text(s)


def extract_venue_from_text(text) -> str:
    if not text or not isinstance(text, str):
        return ""
    m = re.search(r"Venue:\s*(.*?)(?:\s{3,}|$)", text, re.IGNORECASE)
    return m.group(1).strip() if m else ""


def extract_fa(text) -> str:
    if not text or not isinstance(text, str):
        return ""
    t = re.sub(r"^FACULTY\s*ADVISOR\s*:\s*", "", text, flags=re.IGNORECASE).strip()
    t = re.sub(r"\s*Venue:.*$", "", t, flags=re.IGNORECASE).strip()
    return re.sub(r"\s{2,}", " ", t).strip()


def get_inline_venues(ws, r_start, r_end, c_start=1, c_end=22):
    venues = set()
    for row in ws.iter_rows(
        min_row=r_start, max_row=r_end, min_col=c_start, max_col=c_end, values_only=False
    ):
        for c in row:
            if c.value and isinstance(c.value, str):
                v = c.value
                venues.update(
                    m.strip()
                    for m in re.findall(
                        r"\(([^)]*(?:BLOCK|ADMIN|BMS|AD|LH|NA|NEW|MLCP)[^)]*)\)",
                        v,
                        re.IGNORECASE,
                    )
                )
                venues.update(m.strip() for m in re.findall(r"\n\s*\(([^)]+)\)", v))
                venues.update(
                    m.strip()
                    for m in re.findall(r"(?:LH|AD|BMS|NA|ADMIN)\s*\d+", v, re.IGNORECASE)
                )
    return sorted(v for v in venues if v.upper() not in ("LUNCH", "BREAK", "LAB"))


def skip_sheet(name: str) -> bool:
    n = name.strip().lower()
    if "rotation" in n:
        return True
    return any(n.startswith(s) for s in SKIP_SHEETS)


def classify_aiml_ai(sheet_name: str) -> str:
    if "AI " in sheet_name and "AIML" not in sheet_name.upper():
        return "AI"
    return "AIML"


def list_files(data_dir: str, rel: str, exts: tuple[str, ...]) -> list[str]:
    full = os.path.join(data_dir, rel)
    if os.path.isfile(full):
        return [full]
    if not os.path.isdir(full):
        return []
    out = []
    for ext in exts:
        out.extend(glob.glob(os.path.join(full, f"*{ext}")))
    return sorted(
        p
        for p in out
        if not os.path.basename(p).startswith("~$")
        and os.path.basename(p).lower() != "venues.csv"
        and not os.path.basename(p).startswith("_")
    )


def load_csv_records(path: str) -> list[dict]:
    import pandas as pd

    df = pd.read_csv(path)
    df.columns = [c.strip().lower() for c in df.columns]
    recs_map: dict[tuple, dict] = {}
    for _, row in df.iterrows():
        year = str(row.get("year", "")).strip()
        section = str(row.get("section", "")).strip()
        key = (year, section)
        rec = recs_map.setdefault(
            key,
            {
                "year": year,
                "section": section,
                "venue": str(row.get("venue", "") or "(Not specified)").strip()
                or "(Not specified)",
                "fa": str(row.get("faculty", row.get("fa", "")) or "").strip(),
                "timetable": {d: {} for d in DAYS},
                "rotation_venues": [],
            },
        )
        day = str(row.get("day", "")).strip().upper()[:3]
        if day in rec["timetable"]:
            for p in range(1, 9):
                val = row.get(f"p{p}", "")
                rec["timetable"][day][p] = "" if pd.isna(val) else str(val).strip()
        rot = str(row.get("rotation_venues", "") or "").strip()
        if rot:
            rec["rotation_venues"] = [x.strip() for x in rot.split(";") if x.strip()]
    return list(recs_map.values())


def find_day_rows(ws, scan_start=6, scan_end=25) -> dict:
    day_rows = {}
    day_num = {1: "MON", 2: "TUE", 3: "WED", 4: "THU", 5: "FRI"}
    for r in range(scan_start, scan_end + 1):
        av = ws.cell(row=r, column=1).value
        if not av:
            continue
        u = str(av).strip().upper()
        if u.startswith("MON"):
            day_rows["MON"] = r
        elif u.startswith("TUE"):
            day_rows["TUE"] = r
        elif u.startswith("WED"):
            day_rows["WED"] = r
        elif u.startswith("THU"):
            day_rows["THU"] = r
        elif u.startswith("FRI"):
            day_rows["FRI"] = r
        else:
            m = re.match(r"DAY\s*(\d)", u)
            if m:
                dk = day_num.get(int(m.group(1)))
                if dk:
                    day_rows[dk] = r
    return day_rows


def detect_period_cols(ws, day_rows: dict) -> dict:
    if not day_rows:
        return SRM_WIDE_COLS
    sample = day_rows.get("MON") or next(iter(day_rows.values()))
    hits = 0
    for col in SRM_WIDE_COLS:
        v = ws.cell(row=sample, column=col).value
        if v and str(v).strip() and str(v).strip().upper() not in ("BREAK", "LUNCH"):
            hits += 1
    if hits >= 3:
        return SRM_WIDE_COLS
    return {}


def cell_text(ws, r, c) -> str:
    v = ws.cell(row=r, column=c).value
    if v is None:
        return ""
    s = str(v).strip()
    if s.upper() in ("BREAK", "LUNCH", "LUNCH BREAK", "0.0"):
        return ""
    return s


def read_tt(ws, day_rows: dict, period_cols: dict | None) -> dict:
    tt = {}
    if period_cols:
        for day, r in day_rows.items():
            tt[day] = {p: cell_text(ws, r, col) for col, p in period_cols.items()}
        return tt
    for day, r in day_rows.items():
        tt[day] = {}
        pnum = 1
        for col in range(2, min(ws.max_column + 1, 16)):
            cv = cell_text(ws, r, col)
            raw = ws.cell(row=r, column=col).value
            if raw and isinstance(raw, str) and raw.strip().upper() in ("BREAK", "LUNCH"):
                continue
            if not cv:
                continue
            tt[day][pnum] = cv
            pnum += 1
            if pnum > 8:
                break
    return tt


def scan_header_meta(ws, max_row=12, max_col=22) -> tuple[str, str, str]:
    venue, fa, section = "", "", ""
    for r in range(1, max_row + 1):
        for c in range(1, max_col + 1):
            v = ws.cell(row=r, column=c).value
            if not v or not isinstance(v, str):
                continue
            if not venue:
                m = re.search(
                    r"(?:Venue|Room\s*No\.?|ROOM\s*NO)\s*:?\s*(.+)$",
                    v,
                    re.IGNORECASE | re.DOTALL,
                )
                if m and m.group(1).strip():
                    venue = re.sub(r"\s+", " ", m.group(1).strip())
            if not fa:
                low = v.lower()
                if "faculty" in low or "advisor" in low or "incharge" in low or "in-charge" in low:
                    fa = extract_fa(v)
                    if not fa:
                        fa = re.sub(
                            r"^.*?(?:ADVISOR|INCHARGE|IN-CHARGE)\s*:?\s*",
                            "",
                            v,
                            flags=re.IGNORECASE,
                        ).strip()
                        fa = re.sub(r"\s*Venue:.*$", "", fa, flags=re.IGNORECASE).strip()
            if "YEAR/SEM/SEC" in v.upper() and not section:
                sm = re.search(r"YEAR/SEM/SEC\s*:\s*(.*)", v, re.IGNORECASE)
                if sm:
                    section = sm.group(1).strip().rstrip("'\" ")
    return venue, fa, section


def parse_excel_sections(data_dir: str, cfg: dict) -> list[dict]:
    recs = []
    files = list_files(data_dir, cfg["path"], (".xlsx", ".xlsm"))
    sheet_group = cfg.get("sheet_group")
    for fp in files:
        yl_file = year_from_text(os.path.basename(fp))
        try:
            wb = openpyxl.load_workbook(fp, data_only=True)
        except Exception:
            continue
        for sn in wb.sheetnames:
            if skip_sheet(sn):
                continue
            if sheet_group and classify_aiml_ai(sn) != sheet_group:
                continue
            ws = wb[sn]
            venue, fa, header_sec = scan_header_meta(ws)
            yl = yl_file or year_from_info(str(ws["A5"].value or "")) or year_from_info(
                str(ws["A6"].value or "")
            ) or year_from_info(str(ws["A7"].value or "")) or year_from_text(sn)
            if not yl:
                yl = sn.strip()
            dr = find_day_rows(ws)
            if not dr:
                continue
            pcols = detect_period_cols(ws, dr)
            tt = read_tt(ws, dr, pcols or None)
            minr = min(dr.values())
            maxr = max(dr.values())
            rot = get_inline_venues(ws, minr, maxr)
            display_venue = venue
            if not display_venue and rot:
                display_venue = f"Rotation: {', '.join(rot[:3])}"
            recs.append(
                {
                    "year": yl,
                    "section": (header_sec or sn).strip(),
                    "venue": display_venue or "(Not specified)",
                    "fa": fa,
                    "timetable": tt,
                    "rotation_venues": rot,
                }
            )
    return recs


def parse_mech(data_dir: str, cfg: dict) -> list[dict]:
    recs = []
    files = list_files(data_dir, cfg["path"], (".xlsx", ".xlsm"))
    if not files:
        return recs
    wb = openpyxl.load_workbook(files[0], data_only=True)
    sheet = "Class Master TT" if "Class Master TT" in wb.sheetnames else wb.sheetnames[0]
    ws = wb[sheet]
    period_cols = {"C": 1, "D": 2, "F": 3, "G": 4, "K": 5, "L": 6, "M": 7, "N": 8}
    day_blocks = {
        "MON": {"I Year": 10, "II Year": 11, "III Year": 12, "IV Year": 13},
        "TUE": {"I Year": 14, "II Year": 15, "III Year": 16, "IV Year": 17},
        "WED": {"I Year": 18, "II Year": 19, "III Year": 20, "IV Year": 21},
        "THU": {"I Year": 22, "II Year": 23, "III Year": 24, "IV Year": 25},
        "FRI": {"I Year": 26, "II Year": 27, "III Year": 28, "IV Year": 29},
    }
    for yr in ["I Year", "II Year", "III Year", "IV Year"]:
        tt = {}
        all_venues = set()
        for day, year_rows in day_blocks.items():
            r = year_rows[yr]
            tt[day] = {}
            for col_letter, pnum in period_cols.items():
                cv = ws[f"{col_letter}{r}"].value
                if cv and isinstance(cv, str):
                    cv = cv.strip()
                    if cv.upper() in ("BREAK", "LUNCH BREAK", "BREAK\n"):
                        cv = ""
                    else:
                        for v in re.findall(r"\(([^)]+)\)", cv):
                            cleaned = re.sub(r"\s*\n\s*", " ", v).strip()
                            if cleaned.upper() not in ("", "BREAK", "LUNCH BREAK", "LAB", "TUTORIAL"):
                                all_venues.add(cleaned)
                elif cv:
                    cv = str(cv).strip()
                else:
                    cv = ""
                tt[day][pnum] = cv
        recs.append(
            {
                "year": yr,
                "section": "MECH A",
                "venue": ", ".join(sorted(all_venues)[:3]) if all_venues else "(Not specified)",
                "fa": "",
                "timetable": tt,
                "rotation_venues": sorted(all_venues),
            }
        )
    return recs


def parse_eee_xls(data_dir: str, cfg: dict) -> list[dict]:
    recs = []
    files = list_files(data_dir, cfg["path"], (".xls", ".xlsx"))
    if not files:
        return recs
    wb = xlrd.open_workbook(files[0])
    for sn in wb.sheet_names():
        ws = wb.sheet_by_name(sn)
        si = str(ws.cell_value(3, 2)) if ws.ncols > 2 else ""
        yl = year_from_info(si)
        rv = str(ws.cell_value(3, 6)) if ws.ncols > 6 else ""
        venue = ""
        m = re.search(r"ROOM\s*NO\s*:\s*(.*)", rv, re.IGNORECASE)
        if m:
            venue = f"Room {m.group(1).strip()}"
        ci = str(ws.cell_value(3, 8)) if ws.ncols > 8 else ""
        fa = re.sub(r"^\s*CLASS\s*INCHARGE:\s*", "", ci, flags=re.IGNORECASE).strip()
        dm = {"MON": 7, "TUE": 8, "WED": 9, "THU": 10, "FRI": 11}
        tt = {}
        rot = set()
        for day, r in dm.items():
            tt[day] = {}
            for pi, col in enumerate(range(3, min(12, ws.ncols)), 1):
                cv = str(ws.cell_value(r, col)).strip() if ws.ncols > col else ""
                if cv in ("0.0", ""):
                    cv = ""
                if cv.upper() == "LUNCH":
                    cv = ""
                vm = re.findall(r"\(([^)]*(?:AD|BMS|BLOCK|LH|NA)[^)]*)\)", cv, re.IGNORECASE)
                rot.update(v.strip() for v in vm)
                tt[day][pi] = cv
        recs.append(
            {
                "year": yl or sn.strip(),
                "section": sn.strip(),
                "venue": venue or "(Not specified)",
                "fa": fa,
                "timetable": tt,
                "rotation_venues": sorted(rot),
            }
        )
    return recs


def parse_iot_csbs(data_dir: str, cfg: dict) -> list[dict]:
    recs = []
    files = list_files(data_dir, cfg["path"], (".xlsx", ".xlsm"))
    if not files:
        return recs
    wb = openpyxl.load_workbook(files[0], data_only=True)
    pcols = {"C": 1, "E": 2, "H": 3, "J": 4, "M": 5, "O": 6, "Q": 7, "S": 8}
    configs = {
        "CSBS_TT": [(6, "CSBS"), (32, "CSBS"), (59, "CSBS"), (91, "CSBS")],
        "IOT_TT": [(8, "IoT"), (36, "IoT"), (64, "IoT"), (94, "IoT")],
    }
    want = cfg.get("sub_dept")
    for sheet, secs in configs.items():
        if sheet not in wb.sheetnames:
            continue
        ws = wb[sheet]
        for sr, dept in secs:
            if want and dept != want:
                continue
            si = str(ws[f"A{sr}"].value or "")
            yl = year_from_info(si)
            ft = str(ws[f"F{sr}"].value or "")
            venue = extract_venue_from_text(ft)
            fa = extract_fa(ft)
            dr = find_day_rows(ws, scan_start=sr + 1, scan_end=sr + 14)
            tt = {}
            letter_cols = {ord(k) - 64: v for k, v in pcols.items()}
            for day, r in dr.items():
                tt[day] = {p: cell_text(ws, r, col) for col, p in letter_cols.items()}
            minr = min(dr.values()) if dr else sr + 4
            maxr = max(dr.values()) if dr else sr + 8
            rot = get_inline_venues(ws, minr, maxr)
            secname = si.split(":")[-1].strip() if ":" in si else dept
            recs.append(
                {
                    "year": yl or si,
                    "section": f"{dept} - {secname}",
                    "venue": venue
                    or ("Rotation: " + ", ".join(rot[:3]) if rot else "(Not specified)"),
                    "fa": fa,
                    "timetable": tt,
                    "rotation_venues": rot,
                    "sub_dept": dept,
                }
            )
    return recs


def parse_ece(data_dir: str, cfg: dict) -> list[dict]:
    recs = []
    files = list_files(data_dir, cfg["path"], (".xlsx", ".xlsm"))
    if not files:
        return recs
    wb = openpyxl.load_workbook(files[0], data_only=True)
    year_map = {"1st ECE": "I Year", "2nd ECE": "II Year", "3rd ECE": "III Year"}
    for sn in wb.sheetnames:
        ws = wb[sn]
        yl = year_map.get(sn, year_from_text(sn) or sn)
        venue, fa, _ = scan_header_meta(ws)
        rv = str(ws.cell(row=4, column=8).value or "")
        m = re.search(r"Room\s*No\.?\s*:?\s*(.*)", rv, re.IGNORECASE)
        if m:
            venue = m.group(1).strip() or venue
        ci = str(ws.cell(row=5, column=5).value or "")
        m2 = re.search(r"CLASS\s*INCHARGE\s*:\s*(.*)", ci, re.IGNORECASE)
        if m2:
            fa = m2.group(1).strip()
        sec_raw = str(ws.cell(row=5, column=1).value or "")
        sec_m = re.search(r"CLASS\s*:\s*(.*)", sec_raw, re.IGNORECASE)
        section = sec_m.group(1).strip() if sec_m else sn
        dr = find_day_rows(ws)
        period_cols = {2: 1, 3: 2, 4: 3, 5: 4, 7: 5, 8: 6, 9: 7, 10: 8}
        tt = read_tt(ws, dr, period_cols)
        recs.append(
            {
                "year": yl,
                "section": f"ECE {section}",
                "venue": venue or "(Not specified)",
                "fa": fa,
                "timetable": tt,
                "rotation_venues": [],
            }
        )
    return recs


def parse_it_docx(data_dir: str, cfg: dict) -> list[dict]:
    from docx import Document

    recs = []
    day_map = {"MON": "MON", "TUE": "TUE", "WED": "WED", "THU": "THU", "FRI": "FRI"}
    files = list_files(data_dir, cfg["path"], (".docx",))
    for fp in files:
        yl = year_from_text(os.path.basename(fp)) or "I Year"
        doc = Document(fp)
        sections_info = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            sm = re.search(r"SECTION:\s*([A-Z])", txt, re.IGNORECASE)
            if sm:
                sections_info.append({"section": f"IT {sm.group(1).upper()}", "venue": "", "fa": ""})
            if "class-in-charge" in txt.lower() and sections_info:
                cm = re.search(r"Class-In-Charge:\s*(.*?)(?:\t|Block/)", txt, re.IGNORECASE)
                if cm:
                    sections_info[-1]["fa"] = cm.group(1).strip().rstrip(",").strip()
                vm = re.search(r"Block/\s*Room\s*No:\s*(.*)", txt, re.IGNORECASE)
                if vm:
                    sections_info[-1]["venue"] = vm.group(1).strip()
        tt_tables = [doc.tables[i] for i in range(0, len(doc.tables), 2)]
        for idx, table in enumerate(tt_tables):
            if idx >= len(sections_info):
                break
            info = sections_info[idx]
            tt = {}
            rot = set()
            for ri in range(1, min(6, len(table.rows))):
                cells = [c.text.strip() for c in table.rows[ri].cells]
                day_raw = cells[0].upper() if cells else ""
                day_key = day_map.get(day_raw)
                if not day_key:
                    continue
                tt[day_key] = {}
                col_period = {1: 1, 2: 2, 3: 3, 4: 4, 6: 5, 7: 6, 8: 7, 9: 8}
                if len(cells) >= 11:
                    col_period = {1: 1, 2: 2, 3: 2, 4: 3, 5: 4, 7: 5, 8: 6, 9: 7, 10: 8}
                for ci, pnum in col_period.items():
                    if ci >= len(cells):
                        continue
                    val = cells[ci]
                    if not val or val in ("-", "--") or "LUNCH" in val.upper():
                        tt[day_key][pnum] = ""
                        continue
                    lines = val.split("\n")
                    venues_in_cell = []
                    for line in lines[1:]:
                        line = line.strip()
                        if line and not line.startswith("<") and not line.startswith("-"):
                            venues_in_cell.append(line)
                    rot.update(
                        v
                        for v in venues_in_cell
                        if re.search(r"(Adm|BMS|Block|B-III|LH|NA|NwAdm)", v, re.IGNORECASE)
                    )
                    display = lines[0].strip()
                    if venues_in_cell:
                        display += f" ({venues_in_cell[0]})"
                    tt[day_key][pnum] = display
            recs.append(
                {
                    "year": yl,
                    "section": info["section"],
                    "venue": info["venue"] or "(Not specified)",
                    "fa": info["fa"],
                    "timetable": tt,
                    "rotation_venues": sorted(rot),
                }
            )
    return recs


def parse_civil_docx(data_dir: str, cfg: dict) -> list[dict]:
    from docx import Document

    recs = []

    def parse_tt_table(table):
        day_map = {
            "MON": "MON",
            "TUE": "TUE",
            "WED": "WED",
            "THUR": "THU",
            "THURS": "THU",
            "THU": "THU",
            "FRI": "FRI",
        }
        tt = {}
        for ri in range(1, min(6, len(table.rows))):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            day_raw = cells[0].upper() if cells else ""
            day = day_map.get(day_raw, "")
            if not day:
                continue
            tt[day] = {}
            pnum = 1
            for ci in range(1, len(cells)):
                val = cells[ci]
                if val.upper() in ("", "LUNCH", "LUNCH BREAK") or "BREAK" in val.upper():
                    continue
                tt[day][pnum] = val
                pnum += 1
        return tt

    def parse_info_para(text):
        room, fa = "", ""
        m = re.search(r"Room:\s*(.*?)(?:\s{2,}|$)", text, re.IGNORECASE)
        if m:
            room = m.group(1).strip().rstrip(",")
        m2 = re.search(r"Class-In-Charge:\s*(.*?)(?:\s{3,}|Room:)", text, re.IGNORECASE)
        if m2:
            fa = m2.group(1).strip().rstrip(",")
        return room, fa

    files = list_files(data_dir, cfg["path"], (".docx",))
    files = [f for f in files if "civil" in os.path.basename(f).lower()]
    for fp in files:
        doc = Document(fp)
        name = os.path.basename(fp).lower()
        if "higher" in name:
            info_paras = []
            for p in doc.paragraphs:
                if "room" in p.text.lower() and "class-in-charge" in p.text.lower():
                    info_paras.append(parse_info_para(p.text))
            year_configs = [(0, "II Year", 0), (2, "III Year", 1), (4, "IV Year", 2)]
            for tt_idx, yr, info_idx in year_configs:
                if tt_idx < len(doc.tables):
                    tt = parse_tt_table(doc.tables[tt_idx])
                    room = info_paras[info_idx][0] if info_idx < len(info_paras) else ""
                    fa = info_paras[info_idx][1] if info_idx < len(info_paras) else ""
                    recs.append(
                        {
                            "year": yr,
                            "section": "CIVIL A",
                            "venue": room or "(Not specified)",
                            "fa": fa,
                            "timetable": tt,
                            "rotation_venues": [],
                        }
                    )
        else:
            room, fa = "", ""
            for p in doc.paragraphs:
                if "room" in p.text.lower():
                    room, fa = parse_info_para(p.text)
                    break
            tt = parse_tt_table(doc.tables[0]) if doc.tables else {}
            recs.append(
                {
                    "year": "I Year",
                    "section": "CIVIL A",
                    "venue": room or "(Not specified)",
                    "fa": fa,
                    "timetable": tt,
                    "rotation_venues": [],
                }
            )
    return recs


def _csgt_spec(text: str) -> str:
    u = text.upper()
    if "DATA SCIENCE" in u:
        return "DS"
    if "GAMING" in u:
        return "GT"
    return "CS"


def _csgt_section_label(spec: str, raw: str) -> str:
    s = re.sub(r"\s+", " ", (raw or "").strip())
    s = s.replace("'", "").strip()
    su = s.upper()
    if spec == "GT":
        if su in {"A", "B", "C", "D", "E"}:
            return f"GT {s}"
        if "GT" in su:
            label = s.replace("CSE GT", "GT").replace("CSE-GT", "GT")
            return label if re.search(r"GT\s*[A-E]", label, re.I) else "GT A"
        return "GT A"
    if spec == "DS":
        return s if su.startswith("DS") else f"DS {s}"
    if spec == "CS":
        if su.startswith("CS"):
            return s
        return f"CS {s}"
    return s


def parse_csgt_pdf(data_dir: str, cfg: dict) -> list[dict]:
    """Parse CS / GT / DS odd-semester PDF timetables (one section per page)."""
    import pdfplumber

    files = list_files(data_dir, cfg["path"], (".pdf",))
    by_key: dict[tuple, dict] = {}
    skip = {"BREAK", "LUNCH", "B", "R", "E", "A", "K", "U", "N", "C", "H", "L", "P", "I", "O"}

    for fp in files:
        yl_file = year_from_text(os.path.basename(fp))
        with pdfplumber.open(fp) as doc:
            for page in doc.pages:
                text = page.extract_text() or ""
                if "YEAR/SEM/SEC" not in text.upper():
                    continue
                spec = _csgt_spec(text)
                header = ""
                for line in text.splitlines():
                    if "YEAR/SEM/SEC" in line.upper():
                        header = line
                        break
                ym = re.search(
                    r"YEAR/SEM/SEC\s*:\s*(I{1,3}V?)\s*/\s*[^/]*/\s*(.*?)(?:\s+FACULTY|$)",
                    header,
                    re.IGNORECASE,
                )
                year = {
                    "I": "I Year",
                    "II": "II Year",
                    "III": "III Year",
                    "IV": "IV Year",
                }.get((ym.group(1).upper() if ym else ""), yl_file or "")
                raw_sec = ym.group(2).strip() if ym else "Section"
                section = _csgt_section_label(spec, raw_sec)

                fa = ""
                fm = re.search(r"FACULTY\s*ADVISOR\s*:\s*(.*?)\s*Venue", header, re.IGNORECASE)
                if fm:
                    fa = fm.group(1).strip(" :")
                else:
                    fm = re.search(r"FACULTY\s*ADVISOR\s*:\s*(.*)$", header, re.IGNORECASE)
                    if fm:
                        fa = re.sub(r"\s*Venue.*$", "", fm.group(1), flags=re.IGNORECASE).strip(" :")

                venue = ""
                vm = re.search(r"Venue\s*:?\s*(.+)$", header, re.IGNORECASE)
                if vm:
                    venue = re.sub(r"\s+", " ", vm.group(1)).strip(" :")
                    venue = re.sub(r"\s*/\s*", "/", venue)

                tt = {d: {} for d in DAYS}
                rot = set()
                tables = page.extract_tables() or []
                if tables:
                    table = max(tables, key=lambda t: len(t))
                    period_cols = {}
                    for row in table:
                        if not row:
                            continue
                        first = str(row[0] or "").strip().upper()
                        if first.startswith("PERIOD"):
                            for i, cell in enumerate(row[1:], 1):
                                c = str(cell or "").strip()
                                if c.isdigit() and 1 <= int(c) <= 8:
                                    period_cols[i] = int(c)
                            break
                    if not period_cols:
                        period_cols = {1: 1, 2: 2, 4: 3, 5: 4, 8: 5, 10: 6, 11: 7, 12: 8}
                    for row in table:
                        if not row:
                            continue
                        first = str(row[0] or "").strip().upper()
                        day = None
                        if first.startswith("MON"):
                            day = "MON"
                        elif first.startswith("TUE"):
                            day = "TUE"
                        elif first.startswith("WED"):
                            day = "WED"
                        elif first.startswith("THU"):
                            day = "THU"
                        elif first.startswith("FRI"):
                            day = "FRI"
                        if not day:
                            continue
                        for i, pnum in period_cols.items():
                            if i >= len(row):
                                continue
                            val = re.sub(r"\s+", " ", str(row[i] or "")).strip()
                            if not val or val.upper() in skip or val.upper() in ("BREAK", "LUNCH"):
                                continue
                            if re.fullmatch(r"[BREAKLUNCHIOP]+", val.upper().replace(" ", "")):
                                continue
                            rooms = re.findall(
                                r"\(([^)]*(?:ADMIN|BMS|NA|NEW|BLOCK|AD\s*\d)[^)]*)\)",
                                val,
                                re.IGNORECASE,
                            )
                            rot.update(re.sub(r"\s+", " ", r).strip() for r in rooms)
                            tt[day][pnum] = val

                rec = {
                    "year": year or "I Year",
                    "section": section,
                    "venue": venue or "(Not specified)",
                    "fa": fa,
                    "timetable": tt,
                    "rotation_venues": sorted(rot),
                    "spec": spec,
                }
                key = (rec["year"], spec, re.sub(r"\s+", " ", section.upper()))
                prev = by_key.get(key)
                if not prev or len(venue) >= len(prev.get("venue") or ""):
                    by_key[key] = rec

    order = {"CS": 0, "GT": 1, "DS": 2}
    recs = list(by_key.values())
    recs.sort(key=lambda r: (year_sort(r["year"]), order.get(r.get("spec", ""), 9), r["section"]))
    return recs


PARSERS = {
    "excel_sections": parse_excel_sections,
    "mech": parse_mech,
    "eee_xls": parse_eee_xls,
    "iot_csbs": parse_iot_csbs,
    "ece": parse_ece,
    "it_docx": parse_it_docx,
    "civil_docx": parse_civil_docx,
    "csgt_pdf": parse_csgt_pdf,
}


def load_department(data_dir: str, cfg: dict) -> list[dict]:
    # Folder-level venues.csv overrides Excel/Word for that department only.
    folder = os.path.join(data_dir, cfg.get("path", "."))
    if os.path.isdir(folder):
        override = os.path.join(folder, "venues.csv")
        if os.path.isfile(override):
            return load_csv_records(override)
    parser = PARSERS.get(cfg.get("parser", "excel_sections"), parse_excel_sections)
    return parser(data_dir, cfg)


def load_all(data_dir: str, departments: list[dict]) -> tuple[dict[str, list[dict]], dict[str, str]]:
    out: dict[str, list[dict]] = {}
    errors: dict[str, str] = {}
    for cfg in departments:
        try:
            out[cfg["id"]] = load_department(data_dir, cfg)
        except Exception as exc:
            out[cfg["id"]] = []
            errors[cfg["id"]] = f"{type(exc).__name__}: {exc}"
    return out, errors


def period_room(rec: dict, subject: str) -> str:
    """Room for one period: inline cell venue if present, else home classroom."""
    if subject:
        found = re.findall(r"\(([^)]+)\)", subject)
        for v in found:
            cleaned = re.sub(r"\s+", " ", v).strip()
            if cleaned and cleaned.upper() not in ("BREAK", "LUNCH", "LAB", "TUTORIAL"):
                return cleaned
    home = (rec.get("venue") or "").strip()
    if home and home != "(Not specified)" and not home.lower().startswith("rotation:"):
        return home
    return home or ""


def occupancy_rows(catalog: dict[str, list[dict]]) -> list[dict]:
    rows = []
    for dept, recs in catalog.items():
        for rec in recs:
            for day, periods in (rec.get("timetable") or {}).items():
                for p, subj in periods.items():
                    if not subj:
                        continue
                    rows.append(
                        {
                            "dept": dept,
                            "year": rec.get("year", ""),
                            "section": rec.get("section", ""),
                            "venue": period_room(rec, subj),
                            "home_venue": rec.get("venue", ""),
                            "day": day,
                            "period": int(p),
                            "subject": subj,
                            "fa": rec.get("fa") or "",
                        }
                    )
    return rows


def search_catalog(catalog: dict[str, list[dict]], query: str) -> dict:
    q = query.strip().lower()
    sections, faculty, rooms = [], [], []
    if not q:
        return {"sections": sections, "faculty": faculty, "rooms": rooms}
    seen_fa = set()
    seen_room = set()
    for dept, recs in catalog.items():
        for rec in recs:
            blob = " ".join(
                [
                    dept,
                    rec.get("year", ""),
                    rec.get("section", ""),
                    rec.get("venue", ""),
                    rec.get("fa", ""),
                    " ".join(rec.get("rotation_venues") or []),
                ]
            ).lower()
            if q in blob:
                sections.append({**rec, "dept": dept})
            fa = (rec.get("fa") or "").strip()
            if fa and q in fa.lower() and fa.lower() not in seen_fa:
                seen_fa.add(fa.lower())
                faculty.append({"fa": fa, "dept": dept, "year": rec.get("year"), "section": rec.get("section"), "venue": rec.get("venue")})
            for room in [rec.get("venue", "")] + list(rec.get("rotation_venues") or []):
                if room and q in room.lower() and room.lower() not in seen_room:
                    seen_room.add(room.lower())
                    rooms.append({"venue": room, "dept": dept, "year": rec.get("year"), "section": rec.get("section")})
    return {"sections": sections, "faculty": faculty, "rooms": rooms}


def data_stamp(root: str, data_dir: str) -> str:
    items = []
    cfg = os.path.join(root, "config.json")
    if os.path.isfile(cfg):
        items.append(f"config.json:{os.path.getmtime(cfg):.0f}")
    for dirpath, _, files in os.walk(data_dir):
        for f in files:
            if f.startswith("~$") or f.startswith("."):
                continue
            p = os.path.join(dirpath, f)
            try:
                items.append(f"{os.path.relpath(p, data_dir)}:{os.path.getmtime(p):.0f}:{os.path.getsize(p)}")
            except OSError:
                pass
    return "|".join(sorted(items))
