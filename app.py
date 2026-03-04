"""
SRM Institute of Science and Technology - Ramapuram Campus
All Programme Venue Dashboard
Flow: Select Programme -> Select Day -> See Venues (Year I to III)
"""

import streamlit as st
import openpyxl
import xlrd
import re
import os

st.set_page_config(
    page_title="SRM Ramapuram - Venue Dashboard",
    page_icon="🏫",
    layout="wide",
    initial_sidebar_state="collapsed",
)

BASE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "data")

# ══════════════════════════════════════
# STYLES
# ══════════════════════════════════════
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800;900&display=swap');
.main .block-container { padding-top: 1rem; max-width: 1500px; }

/* Title */
.title-box {
    background: linear-gradient(135deg, #1a237e, #0d47a1, #01579b);
    color: white; padding: 22px 30px; border-radius: 14px;
    text-align: center; margin-bottom: 25px;
    box-shadow: 0 6px 25px rgba(13,71,161,0.3);
}
.title-box h1 { font-family:'Inter'; font-size:24px; font-weight:800; margin:0; }
.title-box p { font-size:13px; opacity:0.85; margin:5px 0 0 0; }

/* Programme cards */
.prog-card {
    background: linear-gradient(145deg, #ffffff, #f0f4ff);
    border: 2px solid #e3f2fd; border-radius: 16px;
    padding: 28px 15px; text-align: center; cursor: pointer;
    transition: all 0.3s ease;
    box-shadow: 0 3px 12px rgba(0,0,0,0.06);
    min-height: 160px;
    display: flex; flex-direction: column;
    align-items: center; justify-content: center;
}
.prog-card:hover {
    transform: translateY(-5px);
    box-shadow: 0 8px 30px rgba(13,71,161,0.18);
    border-color: #1565c0;
}
.prog-icon { font-size: 48px; margin-bottom: 10px; }
.prog-name {
    font-family:'Inter'; font-size:20px; font-weight:800;
    color: #0d47a1; letter-spacing: -0.3px;
}
.prog-sub { font-size:11px; color:#78909c; margin-top:4px; }
.prog-count { font-size:11px; color:#1565c0; font-weight:600; margin-top:2px; }

/* Day buttons */
div.stButton > button {
    font-family:'Inter'; font-weight:700; font-size:16px;
    border-radius: 12px; border: 2px solid #e0e0e0;
    transition: all 0.2s;
}
div.stButton > button:hover {
    border-color: #1565c0; background: #e3f2fd;
}

/* Year header */
.year-header {
    background: linear-gradient(135deg, #1565c0, #0d47a1);
    color: white; padding: 10px 22px; border-radius: 10px;
    font-family:'Inter'; font-weight:700; font-size:15px;
    margin: 18px 0 8px 0; display: inline-block;
    box-shadow: 0 3px 10px rgba(13,71,161,0.25);
}

/* Venue table */
.vtable {
    width: 100%; border-collapse: separate; border-spacing: 0;
    border-radius: 10px; overflow: hidden;
    box-shadow: 0 2px 10px rgba(0,0,0,0.07);
    margin-bottom: 18px; font-family:'Inter';
}
.vtable thead th {
    background: #37474f; color: white;
    padding: 10px 12px; font-size: 12px; font-weight: 600;
    text-align: center; white-space: nowrap;
}
.vtable thead th:first-child { text-align: left; }
.vtable tbody td {
    padding: 8px 10px; font-size: 12px; text-align: center;
    border-bottom: 1px solid #eceff1; color: #37474f;
}
.vtable tbody td:first-child { text-align: left; font-weight: 600; }
.vtable tbody tr:nth-child(even) { background: #f8fbff; }
.vtable tbody tr:nth-child(odd) { background: #fff; }
.vtable tbody tr:hover { background: #e3f2fd; }
.venue-tag {
    display: inline-block; background: #e8f5e9; color: #1b5e20;
    padding: 3px 10px; border-radius: 6px; font-weight: 700;
    font-size: 12px; white-space: nowrap;
}
.rot-tag {
    display: inline-block; background: #fff3e0; color: #e65100;
    padding: 2px 8px; border-radius: 5px; font-weight: 600;
    font-size: 11px; margin: 1px 2px; white-space: nowrap;
}
.subj-cell { font-size: 11px; color: #455a64; }

/* Dept header */
.dept-banner {
    background: linear-gradient(135deg, #e8eaf6, #c5cae9);
    padding: 16px 25px; border-radius: 12px;
    border-left: 5px solid #1a237e; margin-bottom: 18px;
}
.dept-banner h2 { font-family:'Inter'; color:#1a237e; margin:0; font-size:22px; }
.dept-banner p { color:#546e7a; margin:4px 0 0 0; font-size:12px; }

/* Day colors */
.day-mon { background: #fff3e0 !important; }
.day-tue { background: #e8f5e9 !important; }
.day-wed { background: #e3f2fd !important; }
.day-thu { background: #fce4ec !important; }
.day-fri { background: #f3e5f5 !important; }
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════
# EXTRACTION HELPERS
# ══════════════════════════════════════

def extract_venue_from_text(text):
    if not text or not isinstance(text, str):
        return ""
    m = re.search(r'Venue:\s*(.*?)(?:\s{3,}|$)', text, re.IGNORECASE)
    return m.group(1).strip() if m else ""

def extract_fa(text):
    if not text or not isinstance(text, str):
        return ""
    t = re.sub(r'^FACULTY\s*ADVISOR\s*:\s*', '', text, flags=re.IGNORECASE).strip()
    t = re.sub(r'\s*Venue:.*$', '', t, flags=re.IGNORECASE).strip()
    return re.sub(r'\s{2,}', ' ', t).strip()

def get_inline_venues(ws, r_start, r_end, c_start=1, c_end=22):
    venues = set()
    for row in ws.iter_rows(min_row=r_start, max_row=r_end, min_col=c_start, max_col=c_end, values_only=False):
        for c in row:
            if c.value and isinstance(c.value, str):
                v = c.value
                venues.update(m.strip() for m in re.findall(
                    r'\(([^)]*(?:BLOCK|ADMIN|BMS|AD|LH|NA|NEW|MLCP)[^)]*)\)', v, re.IGNORECASE))
                venues.update(m.strip() for m in re.findall(r'\n\s*\(([^)]+)\)', v))
                venues.update(m.strip() for m in re.findall(r'(?:LH|AD|BMS|NA|ADMIN)\s*\d+', v, re.IGNORECASE))
    return sorted(v for v in venues if v.upper() not in ('LUNCH', 'BREAK', 'LAB'))

def year_from_info(s):
    if not s: return ""
    m = re.search(r'(I{1,3}V?)\s*/', s, re.IGNORECASE)
    if m:
        return {'I':'I Year','II':'II Year','III':'III Year','IV':'IV Year'}.get(m.group(1).upper(), '')
    return ""

def year_sort(y):
    return {'I Year':1,'II Year':2,'III Year':3,'IV Year':4}.get(y, 99)


# ══════════════════════════════════════
# PARSERS - return list of section dicts
# Each dict: year, section, venue, fa, rotation_venues,
#   timetable: {DAY: {period_num: subject_str}}
# ══════════════════════════════════════

DAYS = ['MON', 'TUE', 'WED', 'THU', 'FRI']
PERIOD_TIMES = {
    1: '8:30-9:20', 2: '9:20-10:10', 3: '10:20-11:10', 4: '11:10-12:00',
    5: '12:40-1:30', 6: '1:30-2:20', 7: '2:20-3:05', 8: '3:05-3:50'
}

def _read_tt_wide(ws, day_rows, period_cols=None):
    """Read timetable from wide-format sheet (merged columns)."""
    if period_cols is None:
        period_cols = {'C':1,'E':2,'H':3,'J':4,'M':5,'O':6,'Q':7,'S':8}
    tt = {}
    for day, r in day_rows.items():
        tt[day] = {}
        for col, pnum in period_cols.items():
            cv = ws[f'{col}{r}'].value
            if cv and isinstance(cv, str):
                cv = cv.strip()
                if cv.upper() in ('BREAK', 'LUNCH'):
                    cv = ""
            elif cv:
                cv = str(cv).strip()
            else:
                cv = ""
            tt[day][pnum] = cv
    return tt

def _find_day_rows(ws, scan_start=8, scan_end=20):
    """Auto-detect which rows contain MON-FRI."""
    day_rows = {}
    for r in range(scan_start, scan_end + 1):
        av = ws[f'A{r}'].value
        if av and isinstance(av, str):
            u = av.strip().upper()
            if u.startswith('MON'): day_rows['MON'] = r
            elif u.startswith('TUE'): day_rows['TUE'] = r
            elif u.startswith('WED'): day_rows['WED'] = r
            elif u.startswith('THU'): day_rows['THU'] = r
            elif u.startswith('FRI'): day_rows['FRI'] = r
    return day_rows


@st.cache_data
def parse_cse():
    recs = []
    pcols = {'C':1,'E':2,'H':3,'J':4,'M':5,'O':6,'Q':7,'S':8}
    dr = {'MON':10,'TUE':11,'WED':12,'THU':13,'FRI':14}
    for yl, fn in [("I Year","I YEAR CSE.xlsx"),("II Year","II YEAR CSE.xlsx"),("III Year","III YEAR CSE.xlsx")]:
        fp = os.path.join(BASE_DIR,"CSE",fn)
        if not os.path.exists(fp): continue
        wb = openpyxl.load_workbook(fp)
        for sn in wb.sheetnames:
            if sn.lower().startswith('sheet'): continue
            ws = wb[sn]
            venue = ""
            for c in ['P','Q','R','S','T']:
                v = ws[f'{c}7'].value
                if v and isinstance(v,str) and 'venue' in v.lower():
                    m = re.search(r'VENUE:\s*(.*)', v, re.IGNORECASE)
                    if m: venue = m.group(1).strip()
                    break
            if not venue:
                venue = extract_venue_from_text(ws['F7'].value)
            fa = extract_fa(ws['F7'].value)
            if not fa:
                h7 = ws['H7'].value
                fa = extract_fa(h7 if h7 else '')
            tt = _read_tt_wide(ws, dr, pcols)
            rot = get_inline_venues(ws, 10, 14)
            recs.append({'year':yl,'section':sn.strip(),'venue':venue or '(Not specified)',
                         'fa':fa,'timetable':tt,'rotation_venues':rot})
    return recs

@st.cache_data
def parse_csgt():
    recs = []
    pcols = {'C':1,'E':2,'H':3,'J':4,'M':5,'O':6,'Q':7,'S':8}
    for yl, fn in [("I Year","I year CS&GT.xlsx"),("II Year","II year CS&GT.xlsx"),("III Year","III year CS&GT.xlsx")]:
        fp = os.path.join(BASE_DIR,"CS_GT",fn)
        if not os.path.exists(fp): continue
        wb = openpyxl.load_workbook(fp)
        for sn in wb.sheetnames:
            if sn.lower().startswith('sheet') or 'rotation' in sn.lower(): continue
            ws = wb[sn]
            venue, fa = "", ""
            for cr in ['J7','F7','H7','Q7']:
                v = ws[cr].value
                if v and isinstance(v,str):
                    if 'venue' in v.lower() and not venue:
                        venue = extract_venue_from_text(v)
                    if 'faculty' in v.lower() and not fa:
                        fa = extract_fa(v)
            dr = _find_day_rows(ws)
            tt = _read_tt_wide(ws, dr, pcols)
            rot = get_inline_venues(ws, min(dr.values()) if dr else 10, max(dr.values()) if dr else 14)
            recs.append({'year':yl,'section':sn.strip(),'venue':venue or '(Not specified)',
                         'fa':fa,'timetable':tt,'rotation_venues':rot})
    return recs

@st.cache_data
def parse_biotech():
    recs = []
    fp = os.path.join(BASE_DIR,"BIOTECH TT EVEN SEM 25-26  REVISED 8.1.25.xlsx")
    if not os.path.exists(fp): return recs
    wb = openpyxl.load_workbook(fp)
    for sn in wb.sheetnames:
        ws = wb[sn]
        si = str(ws['A5'].value or '')
        yl = year_from_info(si)
        venue, fa = "", ""
        for r in [6,7,8]:
            for c in ['F','G','H','I','J']:
                v = ws[f'{c}{r}'].value
                if v and isinstance(v,str):
                    if 'venue' in v.lower() and not venue:
                        m = re.search(r'VENUE\s*:\s*(.*)', v, re.IGNORECASE)
                        if m: venue = m.group(1).strip()
                    if 'faculty' in v.lower() and not fa:
                        fa = extract_fa(v)
        dr = _find_day_rows(ws)
        tt = {}
        for day, r in dr.items():
            tt[day] = {}
            pnum = 1
            for col in range(2, min(ws.max_column+1, 15)):
                cv = ws.cell(row=r, column=col).value
                if cv and isinstance(cv,str):
                    cv = cv.strip()
                    if cv.upper() in ('BREAK','LUNCH'): continue
                    tt[day][pnum] = cv; pnum += 1
        minr = min(dr.values()) if dr else 10
        maxr = max(dr.values()) if dr else 14
        rot = get_inline_venues(ws, minr, maxr)
        display_venue = venue
        if not display_venue and rot:
            display_venue = f"Rotation: {', '.join(rot[:3])}"
        recs.append({'year':yl or sn,'section':sn.strip(),'venue':display_venue or '(Not specified)',
                     'fa':fa,'timetable':tt,'rotation_venues':rot})
    return recs

@st.cache_data
def parse_biomed():
    recs = []
    fp = os.path.join(BASE_DIR,"Bio-Med","TT ODD SEM FINAL - 1 Bio med.xlsx")
    if not os.path.exists(fp): return recs
    wb = openpyxl.load_workbook(fp)
    for sn in wb.sheetnames:
        if sn.lower().startswith('sheet'): continue
        ws = wb[sn]
        yl = ""
        for r in [5,6,7]:
            yl = year_from_info(str(ws[f'A{r}'].value or ''))
            if yl: break
        if not yl:
            if 'IV' in sn.upper(): yl='IV Year'
            elif 'III' in sn.upper(): yl='III Year'
            elif 'II' in sn: yl='II Year'
            else: yl='I Year'
        venue, fa = "", ""
        for r in [5,6,7,8]:
            for c in ['F','G','H','I','J']:
                v = ws[f'{c}{r}'].value
                if v and isinstance(v,str):
                    if 'venue' in v.lower() and not venue:
                        m = re.search(r'Venue:\s*(.*?)(?:\s{3,}|$)', v, re.IGNORECASE)
                        if m: venue = m.group(1).strip()
                    if 'faculty' in v.lower() and not fa:
                        fa = extract_fa(v)
        dr = _find_day_rows(ws)
        tt = {}
        for day, r in dr.items():
            tt[day] = {}
            pnum = 1
            for col in range(2, min(ws.max_column+1, 15)):
                cv = ws.cell(row=r, column=col).value
                if cv and isinstance(cv,str):
                    cv = cv.strip()
                    if cv.upper() in ('BREAK','LUNCH'): continue
                    tt[day][pnum] = cv; pnum += 1
        minr = min(dr.values()) if dr else 10
        maxr = max(dr.values()) if dr else 14
        rot = get_inline_venues(ws, minr, maxr)
        dv = venue
        if not dv and rot:
            dv = f"Rotation: {', '.join(rot[:3])}"
        recs.append({'year':yl,'section':sn.strip(),'venue':dv or '(Not specified)',
                     'fa':fa,'timetable':tt,'rotation_venues':rot})
    return recs

@st.cache_data
def parse_eee():
    recs = []
    fp = os.path.join(BASE_DIR,"EEE time table.xls")
    if not os.path.exists(fp): return recs
    wb = xlrd.open_workbook(fp)
    for sn in wb.sheet_names():
        ws = wb.sheet_by_name(sn)
        si = str(ws.cell_value(3,2)) if ws.ncols>2 else ''
        yl = year_from_info(si)
        rv = str(ws.cell_value(3,6)) if ws.ncols>6 else ''
        venue = ""
        m = re.search(r'ROOM\s*NO\s*:\s*(.*)', rv, re.IGNORECASE)
        if m: venue = f"Room {m.group(1).strip()}"
        ci = str(ws.cell_value(3,8)) if ws.ncols>8 else ''
        fa = re.sub(r'^\s*CLASS\s*INCHARGE:\s*','',ci,flags=re.IGNORECASE).strip()

        dm = {'MON':7,'TUE':8,'WED':9,'THU':10,'FRI':11}
        tt = {}
        rot = set()
        for day, r in dm.items():
            tt[day] = {}
            for pi, col in enumerate(range(3, min(12, ws.ncols)), 1):
                cv = str(ws.cell_value(r, col)).strip() if ws.ncols>col else ""
                if cv in ('0.0',''): cv = ""
                if cv.upper() == 'LUNCH': cv = ""
                # Extract inline venue
                vm = re.findall(r'\(([^)]*(?:AD|BMS|BLOCK|LH|NA)[^)]*)\)', cv, re.IGNORECASE)
                rot.update(v.strip() for v in vm)
                tt[day][pi] = cv

        recs.append({'year':yl or sn.strip(),'section':sn.strip(),'venue':venue or '(Not specified)',
                     'fa':fa,'timetable':tt,'rotation_venues':sorted(rot)})
    return recs

@st.cache_data
def parse_iot_csbs():
    recs = []
    fp = os.path.join(BASE_DIR,"IOT, CSBS.xlsx")
    if not os.path.exists(fp): return recs
    wb = openpyxl.load_workbook(fp)
    pcols = {'C':1,'E':2,'H':3,'J':4,'M':5,'O':6,'Q':7,'S':8}
    configs = {
        'CSBS_TT': [(6,'CSBS'),(32,'CSBS'),(59,'CSBS'),(91,'CSBS')],
        'IOT_TT':  [(8,'IoT'),(36,'IoT'),(64,'IoT'),(94,'IoT')],
    }
    for sheet, secs in configs.items():
        if sheet not in wb.sheetnames: continue
        ws = wb[sheet]
        for sr, dept in secs:
            si = str(ws[f'A{sr}'].value or '')
            yl = year_from_info(si)
            ft = str(ws[f'F{sr}'].value or '')
            venue = extract_venue_from_text(ft)
            fa = extract_fa(ft)
            dr = {}
            for r in range(sr+1, sr+15):
                av = ws[f'A{r}'].value
                if av and isinstance(av, str):
                    u = av.strip().upper()
                    if u.startswith('MON'): dr['MON']=r
                    elif u.startswith('TUE'): dr['TUE']=r
                    elif u.startswith('WED'): dr['WED']=r
                    elif u.startswith('THU'): dr['THU']=r
                    elif u.startswith('FRI'): dr['FRI']=r
            tt = _read_tt_wide(ws, dr, pcols)
            minr = min(dr.values()) if dr else sr+4
            maxr = max(dr.values()) if dr else sr+8
            rot = get_inline_venues(ws, minr, maxr)
            secname = si.split(':')[-1].strip() if ':' in si else dept
            recs.append({'year':yl or si,'section':f"{dept} - {secname}",
                         'venue':venue or ('Rotation: '+', '.join(rot[:3]) if rot else '(Not specified)'),
                         'fa':fa,'timetable':tt,'rotation_venues':rot,'sub_dept':dept})
    return recs

@st.cache_data
def parse_mech():
    recs = []
    fp = os.path.join(BASE_DIR, "MECH_DAYWise_Class _TIME TABLE EVEN SEMESTER (2025-26).xlsx")
    if not os.path.exists(fp):
        return recs
    wb = openpyxl.load_workbook(fp)
    ws = wb['Class Master TT']

    # Structure: Day-wise master timetable
    # Column A=Day, B=Year, C-N=Periods (C=P1, D=P2, E=Break, F=P3, G=P4, H=Lunch, K=P5, L=P6, M=P7, N=P8)
    # MON: rows 10-13, TUE: 14-17, WED: 18-21, THU: 22-25, FRI: 26-29
    period_cols = {'C': 1, 'D': 2, 'F': 3, 'G': 4, 'K': 5, 'L': 6, 'M': 7, 'N': 8}

    day_blocks = {
        'MON': {'I Year': 10, 'II Year': 11, 'III Year': 12, 'IV Year': 13},
        'TUE': {'I Year': 14, 'II Year': 15, 'III Year': 16, 'IV Year': 17},
        'WED': {'I Year': 18, 'II Year': 19, 'III Year': 20, 'IV Year': 21},
        'THU': {'I Year': 22, 'II Year': 23, 'III Year': 24, 'IV Year': 25},
        'FRI': {'I Year': 26, 'II Year': 27, 'III Year': 28, 'IV Year': 29},
    }

    # Build per-year records
    year_data = {}
    for yr in ['I Year', 'II Year', 'III Year', 'IV Year']:
        tt = {}
        all_venues = set()
        for day, year_rows in day_blocks.items():
            r = year_rows[yr]
            tt[day] = {}
            for col_letter, pnum in period_cols.items():
                cv = ws[f'{col_letter}{r}'].value
                if cv and isinstance(cv, str):
                    cv = cv.strip()
                    if cv.upper() in ('BREAK', 'LUNCH BREAK', 'BREAK\n'):
                        cv = ""
                    else:
                        # Extract venues from parentheses
                        vms = re.findall(r'\(([^)]+)\)', cv)
                        for v in vms:
                            cleaned = re.sub(r'\s*\n\s*', ' ', v).strip()
                            if cleaned.upper() not in ('', 'BREAK', 'LUNCH BREAK', 'LAB', 'TUTORIAL'):
                                all_venues.add(cleaned)
                elif cv:
                    cv = str(cv).strip()
                else:
                    cv = ""
                tt[day][pnum] = cv

        # Determine primary venue (most common one)
        venue_str = ', '.join(sorted(all_venues)[:3]) if all_venues else '(Not specified)'

        year_data[yr] = {
            'year': yr,
            'section': f'MECH A',
            'venue': venue_str,
            'fa': '',
            'timetable': tt,
            'rotation_venues': sorted(all_venues),
        }

    recs = [year_data[yr] for yr in ['I Year', 'II Year', 'III Year', 'IV Year']
            if yr in year_data]
    return recs

@st.cache_data
def parse_bda_cc():
    recs = []
    pcols = {'C':1,'E':2,'H':3,'J':4,'M':5,'O':6,'Q':7,'S':8}
    for yl, fn in [("I Year","I YEAR BDA_CC.xlsx"),("II Year","II YEAR BDA_CC.xlsx"),
                    ("III Year","III YEAR BDA_CC.xlsx"),("IV Year","IV YEAR BDA_CC.xlsx")]:
        fp = os.path.join(BASE_DIR,"BDA_CC",fn)
        if not os.path.exists(fp): continue
        wb = openpyxl.load_workbook(fp)
        for sn in wb.sheetnames:
            if sn.lower().startswith('sheet') or sn.lower().startswith('copy of'): continue
            ws = wb[sn]
            venue, fa = "", ""
            # I Year: venue/FA at J7 or F7; II-IV Year: at F6 or J6
            for cr in ['J7','F7','H7','F6','J6','H6']:
                v = ws[cr].value
                if v and isinstance(v,str):
                    if 'venue' in v.lower() and not venue:
                        venue = extract_venue_from_text(v)
                    if ('faculty' in v.lower() or 'advisor' in v.lower()) and not fa:
                        fa = extract_fa(v)
            dr = _find_day_rows(ws, scan_start=8, scan_end=20)
            if not dr:
                dr = _find_day_rows(ws, scan_start=5, scan_end=16)
            tt = _read_tt_wide(ws, dr, pcols)
            minr = min(dr.values()) if dr else 10
            maxr = max(dr.values()) if dr else 14
            rot = get_inline_venues(ws, minr, maxr)
            recs.append({'year':yl,'section':sn.strip(),'venue':venue or '(Not specified)',
                         'fa':fa,'timetable':tt,'rotation_venues':rot})
    return recs

@st.cache_data
def parse_civil():
    from docx import Document
    recs = []

    # Helper: parse a docx timetable table (6 rows: header + Mon-Fri)
    def parse_tt_table(table):
        day_map = {'MON':'MON','TUE':'TUE','WED':'WED','THUR':'THU','THURS':'THU','THU':'THU','FRI':'FRI'}
        tt = {}
        for ri in range(1, min(6, len(table.rows))):
            cells = [c.text.strip() for c in table.rows[ri].cells]
            day_raw = cells[0].upper() if cells else ''
            day = day_map.get(day_raw, '')
            if not day:
                continue
            tt[day] = {}
            # Map cell positions to period numbers (skip BREAK and LUNCH columns)
            pnum = 1
            for ci in range(1, len(cells)):
                val = cells[ci]
                if val.upper() in ('', 'LUNCH', 'LUNCH BREAK') or 'BREAK' in val.upper():
                    continue
                tt[day][pnum] = val
                pnum += 1
        return tt

    # Helper: extract room and FA from paragraph text
    def parse_info_para(text):
        room, fa = '', ''
        m = re.search(r'Room:\s*(.*?)(?:\s{2,}|$)', text, re.IGNORECASE)
        if m:
            room = m.group(1).strip().rstrip(',')
        m2 = re.search(r'Class-In-Charge:\s*(.*?)(?:\s{3,}|Room:)', text, re.IGNORECASE)
        if m2:
            fa = m2.group(1).strip().rstrip(',')
        return room, fa

    # ── File 1: I Year (2nd sem) ──
    fp1 = os.path.join(BASE_DIR, "Civil Even second sem TT 2025.docx")
    if os.path.exists(fp1):
        doc = Document(fp1)
        room, fa = '', ''
        for p in doc.paragraphs:
            if 'room' in p.text.lower():
                room, fa = parse_info_para(p.text)
                break
        tt = parse_tt_table(doc.tables[0]) if doc.tables else {}
        recs.append({'year':'I Year','section':'CIVIL A',
                     'venue':room or '(Not specified)','fa':fa,
                     'timetable':tt,'rotation_venues':[]})

    # ── File 2: II, III, IV Year (higher semesters) ──
    fp2 = os.path.join(BASE_DIR, "Civil Even sem TT higher semester-2025.docx")
    if os.path.exists(fp2):
        doc = Document(fp2)
        # Pairs: (timetable_table_idx, year_label)
        # Table 0=II Year, Table 2=III Year, Table 4=IV Year
        # Info paragraphs appear before each timetable
        info_paras = []
        for p in doc.paragraphs:
            if 'room' in p.text.lower() and 'class-in-charge' in p.text.lower():
                room, fa = parse_info_para(p.text)
                info_paras.append((room, fa))

        year_configs = [
            (0, 'II Year', 0),   # table 0, info_paras[0]
            (2, 'III Year', 1),  # table 2, info_paras[1]
            (4, 'IV Year', 2),   # table 4, info_paras[2]
        ]
        for tt_idx, yr, info_idx in year_configs:
            if tt_idx < len(doc.tables):
                tt = parse_tt_table(doc.tables[tt_idx])
                room = info_paras[info_idx][0] if info_idx < len(info_paras) else ''
                fa = info_paras[info_idx][1] if info_idx < len(info_paras) else ''
                recs.append({'year':yr,'section':'CIVIL A',
                             'venue':room or '(Not specified)','fa':fa,
                             'timetable':tt,'rotation_venues':[]})

    return recs


# ══════════════════════════════════════
# DEPARTMENT REGISTRY
# ══════════════════════════════════════

DEPARTMENTS = {
    "CSE":      {"name":"Computer Science & Engineering",       "icon":"💻","parser":parse_cse},
    "CS & GT":  {"name":"Computer Science & Gaming Technology", "icon":"🎮","parser":parse_csgt},
    "BIOTECH":  {"name":"Biotechnology",                        "icon":"🧬","parser":parse_biotech},
    "Bio-Med":  {"name":"Biomedical Engineering",               "icon":"🏥","parser":parse_biomed},
    "MECH":     {"name":"Mechanical Engineering",                "icon":"⚙️","parser":parse_mech},
    "CIVIL":    {"name":"Civil Engineering",                    "icon":"🏗️","parser":parse_civil},
    "EEE":      {"name":"Electrical & Electronics Engineering", "icon":"⚡","parser":parse_eee},
    "CSBS":     {"name":"Computer Science & Business Systems",  "icon":"📊",
                 "parser": lambda: [r for r in parse_iot_csbs() if r.get('sub_dept')=='CSBS']},
    "IoT":      {"name":"Internet of Things",                   "icon":"📡",
                 "parser": lambda: [r for r in parse_iot_csbs() if r.get('sub_dept')=='IoT']},
    "ECE":      {"name":"Electronics & Communication Engg",     "icon":"📟",
                 "parser": lambda: []},
    "BDA & CC": {"name":"Big Data Analytics & Cloud Computing", "icon":"📈",
                 "parser": parse_bda_cc},
}

DAY_NAMES = {'MON':'Monday','TUE':'Tuesday','WED':'Wednesday','THU':'Thursday','FRI':'Friday'}
DAY_COLORS = {'MON':'#FF9800','TUE':'#4CAF50','WED':'#2196F3','THU':'#E91E63','FRI':'#9C27B0'}
DAY_BG = {'MON':'#fff3e0','TUE':'#e8f5e9','WED':'#e3f2fd','THU':'#fce4ec','FRI':'#f3e5f5'}


# ══════════════════════════════════════
# RENDER: Day-wise venue view
# ══════════════════════════════════════

def render_day_venue(records, day_key):
    """Show all sections for a specific day, grouped by year, with venue and subjects per period."""
    years = sorted(set(r['year'] for r in records), key=year_sort)
    day_name = DAY_NAMES[day_key]
    bg = DAY_BG[day_key]
    clr = DAY_COLORS[day_key]

    st.markdown(f"""
    <div style="background:{bg}; border-left:5px solid {clr}; padding:12px 20px;
                border-radius:10px; margin-bottom:15px;">
        <span style="font-family:'Inter'; font-size:20px; font-weight:800; color:{clr};">
            {day_name}
        </span>
        <span style="font-size:13px; color:#546e7a; margin-left:10px;">
            Venue & Subject details for all sections
        </span>
    </div>
    """, unsafe_allow_html=True)

    for yr in years:
        yr_recs = [r for r in records if r['year'] == yr]
        if not yr_recs:
            continue

        st.markdown(f'<div class="year-header">{yr} &mdash; {len(yr_recs)} Sections</div>',
                    unsafe_allow_html=True)

        # Build table
        html = '<table class="vtable"><thead><tr>'
        html += '<th style="min-width:80px">Section</th>'
        html += '<th style="min-width:130px">Venue</th>'
        for p in range(1, 9):
            html += f'<th>P{p}<br><span style="font-size:10px;font-weight:400">{PERIOD_TIMES[p]}</span></th>'
        html += '<th>Rotation Venues</th>'
        html += '</tr></thead><tbody>'

        for rec in yr_recs:
            tt_day = rec['timetable'].get(day_key, {})
            html += '<tr>'
            html += f'<td style="font-weight:700">{rec["section"]}</td>'
            html += f'<td><span class="venue-tag">{rec["venue"]}</span></td>'

            for p in range(1, 9):
                subj = tt_day.get(p, '')
                if not subj:
                    html += '<td style="color:#bdbdbd">-</td>'
                else:
                    # Clean subject: truncate and show
                    display = subj.replace('\n', ' ')
                    if len(display) > 30:
                        display = display[:28] + '..'
                    html += f'<td class="subj-cell" title="{subj}">{display}</td>'

            # Rotation venues
            if rec.get('rotation_venues'):
                rot_html = ' '.join(f'<span class="rot-tag">{v}</span>' for v in rec['rotation_venues'])
                html += f'<td>{rot_html}</td>'
            else:
                html += '<td style="color:#bdbdbd">-</td>'

            html += '</tr>'

        html += '</tbody></table>'
        st.markdown(html, unsafe_allow_html=True)


# ══════════════════════════════════════
# MAIN APP
# ══════════════════════════════════════

def main():
    if 'dept' not in st.session_state:
        st.session_state.dept = None
    if 'day' not in st.session_state:
        st.session_state.day = None

    # ── Title ──
    st.markdown("""
    <div class="title-box">
        <h1>SRM Institute of Science and Technology &mdash; Ramapuram Campus</h1>
        <p>All Programme Venue Dashboard &nbsp;|&nbsp; 2025-26 Even Semester</p>
    </div>
    """, unsafe_allow_html=True)

    # ════════════════════════════════
    # SCREEN 1: Select Programme
    # ════════════════════════════════
    if st.session_state.dept is None:
        dept_keys = list(DEPARTMENTS.keys())
        # Row 1: first 4 departments
        cols = st.columns(4)
        for i in range(min(4, len(dept_keys))):
            dk = dept_keys[i]
            d = DEPARTMENTS[dk]
            with cols[i]:
                if st.button(f"{d['icon']}\n\n{dk}", key=f"d_{dk}", use_container_width=True,
                             help=d['name']):
                    st.session_state.dept = dk
                    st.session_state.day = None
                    st.rerun()
                st.caption(d['name'])

        # Row 2: next 4+ departments
        if len(dept_keys) > 4:
            cols2 = st.columns(4)
            for i in range(4, min(8, len(dept_keys))):
                dk = dept_keys[i]
                d = DEPARTMENTS[dk]
                with cols2[i - 4]:
                    if st.button(f"{d['icon']}\n\n{dk}", key=f"d_{dk}", use_container_width=True,
                                 help=d['name']):
                        st.session_state.dept = dk
                        st.session_state.day = None
                        st.rerun()
                    st.caption(d['name'])

        # Row 3: remaining departments (9+)
        if len(dept_keys) > 8:
            cols3 = st.columns(4)
            for i in range(8, len(dept_keys)):
                dk = dept_keys[i]
                d = DEPARTMENTS[dk]
                with cols3[i - 8]:
                    if st.button(f"{d['icon']}\n\n{dk}", key=f"d_{dk}", use_container_width=True,
                                 help=d['name']):
                        st.session_state.dept = dk
                        st.session_state.day = None
                        st.rerun()
                    st.caption(d['name'])

        st.markdown("")
        st.markdown("### 👆 Select a Programme")
        st.markdown("*Click on any department above to view venue details from I Year to III Year*")

        return

    # ════════════════════════════════
    # SCREEN 2: Select Day
    # ════════════════════════════════
    dept_key = st.session_state.dept
    dept = DEPARTMENTS[dept_key]
    records = dept['parser']()

    if st.session_state.day is None:
        # Back button
        if st.button("← Back to All Programmes", key="back_dept"):
            st.session_state.dept = None
            st.rerun()

        st.markdown(f"""
        <div class="dept-banner">
            <h2>{dept['icon']} {dept_key} &mdash; {dept['name']}</h2>
            <p>{len(records)} sections across {len(set(r['year'] for r in records))} years &nbsp;|&nbsp;
               Select a day to see venues from I Year to III Year</p>
        </div>
        """, unsafe_allow_html=True)

        if not records:
            st.warning("ECE timetables are in PDF format and cannot be parsed automatically. "
                       "Please provide Excel files for ECE to include them.")
            return

        st.markdown("### 📅 Select a Day")
        st.markdown("")

        day_cols = st.columns(5)
        for i, (dk, dn) in enumerate(DAY_NAMES.items()):
            with day_cols[i]:
                clr = DAY_COLORS[dk]
                if st.button(f"📌\n\n{dn}", key=f"day_{dk}", use_container_width=True):
                    st.session_state.day = dk
                    st.rerun()

        # Also show quick venue summary
        st.markdown("---")
        st.markdown("### 📍 Quick Venue Summary (All Days)")
        years = sorted(set(r['year'] for r in records), key=year_sort)
        for yr in years:
            yr_recs = [r for r in records if r['year'] == yr]
            st.markdown(f'<div class="year-header">{yr}</div>', unsafe_allow_html=True)
            html = '<table class="vtable"><thead><tr>'
            html += '<th>Section</th><th>Venue</th><th>Faculty Advisor</th><th>Rotation Venues</th>'
            html += '</tr></thead><tbody>'
            for rec in yr_recs:
                rot = ' '.join(f'<span class="rot-tag">{v}</span>' for v in rec.get('rotation_venues',[])) or '-'
                html += f"""<tr>
                    <td style="font-weight:700">{rec['section']}</td>
                    <td><span class="venue-tag">{rec['venue']}</span></td>
                    <td>{rec['fa']}</td>
                    <td>{rot}</td>
                </tr>"""
            html += '</tbody></table>'
            st.markdown(html, unsafe_allow_html=True)

        return

    # ════════════════════════════════
    # SCREEN 3: Show Day-wise Venues
    # ════════════════════════════════
    col1, col2 = st.columns([1, 1])
    with col1:
        if st.button("← Back to Days", key="back_day"):
            st.session_state.day = None
            st.rerun()
    with col2:
        if st.button("← Back to All Programmes", key="back_all"):
            st.session_state.dept = None
            st.session_state.day = None
            st.rerun()

    st.markdown(f"""
    <div class="dept-banner">
        <h2>{dept['icon']} {dept_key} &mdash; {DAY_NAMES[st.session_state.day]}</h2>
        <p>Venue details for all sections from I Year to III Year</p>
    </div>
    """, unsafe_allow_html=True)

    render_day_venue(records, st.session_state.day)

    # Option to view other days without going back
    st.markdown("---")
    st.markdown("**Switch Day:**")
    other_cols = st.columns(5)
    for i, (dk, dn) in enumerate(DAY_NAMES.items()):
        with other_cols[i]:
            disabled = dk == st.session_state.day
            if st.button(dn, key=f"switch_{dk}", use_container_width=True, disabled=disabled):
                st.session_state.day = dk
                st.rerun()


if __name__ == "__main__":
    main()
